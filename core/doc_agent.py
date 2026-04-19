"""
DocAgent — 書類作成エージェント（提案書・企画書・報告書など）
Sprint 2 Feature L
"""
from __future__ import annotations

import json
import logging
import re
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Callable

logger = logging.getLogger(__name__)


@dataclass
class DocResult:
    success: bool
    file_path: str
    doc_type: str
    title: str
    sections: list[str]
    message: str


# 書類種別の判定キーワードマッピング
_DOC_TYPE_KEYWORDS: list[tuple[str, list[str]]] = [
    ("提案書",    ["提案書"]),
    ("企画書",    ["企画書"]),
    ("報告書",    ["報告書"]),
    ("議事録",    ["議事録"]),
    ("メール",    ["メール"]),
    ("資料",      ["資料", "書類"]),
]

# 書類種別ごとの典型セクション構成
_DOC_SECTIONS: dict[str, list[str]] = {
    "提案書": ["はじめに", "現状と課題", "提案内容", "期待効果", "実施スケジュール", "まとめ"],
    "企画書": ["企画概要", "背景・目的", "ターゲット", "実施内容", "予算・リソース", "スケジュール", "まとめ"],
    "報告書": ["概要", "実施内容", "結果・成果", "課題と考察", "今後の方針"],
    "議事録": ["開催概要", "出席者", "議題", "討議内容", "決定事項", "次回予定"],
    "メール": ["件名", "本文", "結び"],
    "資料":   ["はじめに", "内容", "まとめ"],
}

# can_handle 用キーワード
_HANDLE_KEYWORDS = ["提案書", "企画書", "報告書", "議事録", "資料", "書類"]


class DocAgent:
    """自然言語の依頼から Word ドキュメントを作成するエージェント。"""

    def __init__(self, base_dir: Path, llm_fn: Callable[[str], str]) -> None:
        self.base_dir = Path(base_dir)
        self.llm_fn = llm_fn
        self._output_dir = self.base_dir / "data" / "documents"
        self._patterns_path = self.base_dir / "data" / "doc_agent_patterns.json"
        self._output_dir.mkdir(parents=True, exist_ok=True)

    # ─────────────────────────────────────────────────────────
    # 公開 API
    # ─────────────────────────────────────────────────────────

    def can_handle(self, user_input: str) -> bool:
        """書類作成依頼かどうかを判定する。"""
        return any(kw in user_input for kw in _HANDLE_KEYWORDS)

    def create(self, user_request: str) -> DocResult:
        """書類を作成して DocResult を返す。"""
        doc_type = self._detect_doc_type(user_request)
        title = self._extract_title(user_request, doc_type)
        sections = _DOC_SECTIONS.get(doc_type, _DOC_SECTIONS["資料"])

        try:
            # セクションコンテンツを生成
            section_contents: dict[str, str] = {}
            for section in sections:
                content = self._generate_section(user_request, doc_type, title, section)
                section_contents[section] = content

            # Word ファイルに出力
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            safe_title = re.sub(r'[\\/:*?"<>|]', "_", title)[:50]
            filename = f"{timestamp}_{safe_title}.docx"
            file_path = self._output_dir / filename

            self._write_docx(file_path, doc_type, title, sections, section_contents)

            # 成功パターンを保存
            self._save_pattern(user_request, doc_type, title, len(sections))

            message = (
                f"✅ {doc_type}を作成したよ！\n"
                f"📄 タイトル: {title}\n"
                f"📁 保存先: {file_path}\n"
                f"📑 セクション数: {len(sections)}"
            )
            return DocResult(
                success=True,
                file_path=str(file_path),
                doc_type=doc_type,
                title=title,
                sections=sections,
                message=message,
            )

        except Exception as exc:
            logger.warning("[DocAgent] ドキュメント作成失敗: %s", exc)
            return DocResult(
                success=False,
                file_path="",
                doc_type=doc_type,
                title=title,
                sections=sections,
                message=f"書類の作成中にエラーが発生したよ: {exc}",
            )

    # ─────────────────────────────────────────────────────────
    # プライベートメソッド
    # ─────────────────────────────────────────────────────────

    def _detect_doc_type(self, user_request: str) -> str:
        """ユーザーリクエストから書類種別を判定する。"""
        for doc_type, keywords in _DOC_TYPE_KEYWORDS:
            if any(kw in user_request for kw in keywords):
                return doc_type
        return "資料"

    def _extract_title(self, user_request: str, doc_type: str) -> str:
        """書類タイトルをリクエストから抽出する。"""
        # 「○○の提案書」「○○についての企画書」などのパターンを検出
        patterns = [
            re.compile(r"(.+?)(?:の|に関する|についての|についての?)" + re.escape(doc_type)),
            re.compile(r"(.+?)(?:を|の)" + re.escape(doc_type)),
        ]
        for pat in patterns:
            m = pat.search(user_request)
            if m:
                candidate = m.group(1).strip()
                if candidate and len(candidate) < 50:
                    return f"{candidate}の{doc_type}"

        # フォールバック: LLM にタイトルを生成させる
        try:
            prompt = (
                f"以下の依頼に合う{doc_type}のタイトルを15文字以内で答えてください。\n"
                f"依頼: {user_request}\n"
                "タイトルのみを出力してください。"
            )
            raw = self.llm_fn(prompt)
            title = raw.strip().splitlines()[0].strip()
            return title if title else f"{doc_type}"
        except Exception:
            return f"{doc_type}"

    def _generate_section(
        self,
        user_request: str,
        doc_type: str,
        title: str,
        section_name: str,
    ) -> str:
        """1セクションのコンテンツを LLM で生成する。"""
        prompt = (
            f"書類: {doc_type}「{title}」\n"
            f"依頼内容: {user_request}\n"
            f"セクション「{section_name}」の内容を日本語で書いてください。\n"
            "200〜400字程度で、箇条書きまたは段落形式で書いてください。"
        )
        try:
            return self.llm_fn(prompt).strip()
        except Exception as exc:
            logger.warning("[DocAgent] セクション生成失敗 (%s): %s", section_name, exc)
            return f"（{section_name}の内容を記載してください）"

    def _write_docx(
        self,
        file_path: Path,
        doc_type: str,
        title: str,
        sections: list[str],
        section_contents: dict[str, str],
    ) -> None:
        """python-docx を使って Word ファイルを書き出す。"""
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # タイトル
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 書類種別・日付
        meta = doc.add_paragraph()
        meta.add_run(f"{doc_type}　作成日: {datetime.now().strftime('%Y年%m月%d日')}")
        meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        doc.add_paragraph()  # 空行

        # 各セクション
        for section in sections:
            doc.add_heading(section, level=1)
            content = section_contents.get(section, "")
            if content:
                # 箇条書き行とそれ以外を分けて処理
                for line in content.splitlines():
                    stripped = line.strip()
                    if not stripped:
                        continue
                    if stripped.startswith(("-", "・", "•", "*")):
                        item_text = re.sub(r"^[-・•\*]\s*", "", stripped)
                        p = doc.add_paragraph(item_text, style="List Bullet")
                    else:
                        doc.add_paragraph(stripped)

        doc.save(str(file_path))

    def _save_pattern(
        self,
        user_request: str,
        doc_type: str,
        title: str,
        section_count: int,
    ) -> None:
        """成功パターンを JSON に追記保存する。"""
        try:
            patterns: list[dict] = []
            if self._patterns_path.exists():
                with open(self._patterns_path, encoding="utf-8") as f:
                    patterns = json.load(f)
        except Exception:
            patterns = []

        patterns.append(
            {
                "timestamp": datetime.now().isoformat(),
                "request": user_request[:200],
                "doc_type": doc_type,
                "title": title,
                "section_count": section_count,
            }
        )
        patterns = patterns[-100:]
        try:
            with open(self._patterns_path, "w", encoding="utf-8") as f:
                json.dump(patterns, f, ensure_ascii=False, indent=2)
        except Exception as exc:
            logger.warning("[DocAgent] パターン保存失敗: %s", exc)
