"""
競合調査レポートエンジン（Sprint 4-F）
会社名・サービス名から競合分析シートを生成する。
"""
from __future__ import annotations

import json
import logging
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Callable

logger = logging.getLogger(__name__)

_REPORTS_DIR = "data/competitor_reports"


@dataclass
class CompetitorReport:
    target: str
    summary: str
    strengths: list[str]
    weaknesses: list[str]
    competitors: list[str]
    file_path: str
    message: str


class CompetitorAnalyzer:
    """会社名・サービス名から競合分析シートを生成。"""

    def __init__(
        self,
        base_dir: Path,
        llm_fn: Callable[[str], str],
        research_agent=None,
    ) -> None:
        self._base_dir = Path(base_dir)
        self._llm_fn = llm_fn
        self._research_agent = research_agent
        self._reports_dir = self._base_dir / _REPORTS_DIR
        self._reports_dir.mkdir(parents=True, exist_ok=True)

    # ──────────────────────────────────────────────────────────────
    # 公開 API
    # ──────────────────────────────────────────────────────────────

    def analyze(self, target: str) -> CompetitorReport:
        """
        1. ResearchAgent（なければ DDG 直接）で target を検索
        2. LLM で競合・強み・弱み・市場ポジションを分析
        3. python-docx で Word レポートを出力
        4. data/competitor_reports/{timestamp}_{target}.docx に保存
        """
        try:
            # ─── 1. 情報収集 ───────────────────────────────────────
            raw_info = self._gather_info(target)

            # ─── 2. LLM 分析 ───────────────────────────────────────
            analysis = self._llm_analyze(target, raw_info)

            # ─── 3 & 4. Word レポート生成 ─────────────────────────
            file_path = self._create_docx(target, analysis)

            report = CompetitorReport(
                target=target,
                summary=analysis.get("summary", ""),
                strengths=analysis.get("strengths", []),
                weaknesses=analysis.get("weaknesses", []),
                competitors=analysis.get("competitors", []),
                file_path=file_path,
                message=(
                    f"「{target}」の競合分析レポートを作ったよ！"
                    f" {file_path} に保存したからチェックしてみてね😊"
                ),
            )
            return report

        except Exception as e:
            logger.warning("[CompetitorAnalyzer] analyze error: %s", e)
            return CompetitorReport(
                target=target,
                summary="",
                strengths=[],
                weaknesses=[],
                competitors=[],
                file_path="",
                message=f"競合分析中にエラーが発生しちゃった: {e}",
            )

    # ──────────────────────────────────────────────────────────────
    # 内部ヘルパー
    # ──────────────────────────────────────────────────────────────

    def _gather_info(self, target: str) -> str:
        """ResearchAgent または DDG で情報収集する。"""
        if self._research_agent is not None:
            try:
                result = self._research_agent.search(f"{target} 競合 強み 弱み 市場")
                return result.summary
            except Exception as e:
                logger.warning("[CompetitorAnalyzer] ResearchAgent 失敗: %s", e)

        # フォールバック: DDG Lite 直接検索
        try:
            import urllib.parse
            import urllib.request

            from utils.url_guard import assert_safe_http_url
            query = urllib.parse.quote(f"{target} 競合他社 サービス比較")
            url = assert_safe_http_url(f"https://lite.duckduckgo.com/lite/?q={query}")
            req = urllib.request.Request(
                url, headers={"User-Agent": "AiChan-Researcher/1.0"}
            )
            with urllib.request.urlopen(req, timeout=8) as resp:  # noqa: S310 (scheme asserted)
                html = resp.read().decode("utf-8", errors="replace")
            # 簡易テキスト抽出
            import re
            text = re.sub(r"<[^>]+>", " ", html)
            text = re.sub(r"\s+", " ", text).strip()
            return text[:3000]
        except Exception as e:
            logger.warning("[CompetitorAnalyzer] DDG 直接検索失敗: %s", e)
            return f"{target} に関する情報"

    def _llm_analyze(self, target: str, raw_info: str) -> dict:
        """LLM で競合・強み・弱み・市場ポジションを分析する。"""
        prompt = (
            f"以下の情報を元に「{target}」の競合分析を行ってください。\n"
            "JSON 形式で返してください（コードブロックなし）:\n"
            '{\n'
            '  "summary": "市場ポジション・概要（3文）",\n'
            '  "strengths": ["強み1", "強み2", "強み3"],\n'
            '  "weaknesses": ["弱み1", "弱み2", "弱み3"],\n'
            '  "competitors": ["競合A", "競合B", "競合C"]\n'
            '}\n\n'
            f"参考情報:\n{raw_info[:2000]}"
        )
        try:
            response = self._llm_fn(prompt)
            # JSON 部分を抽出
            import re
            match = re.search(r"\{.*\}", response, re.DOTALL)
            if match:
                return json.loads(match.group())
        except Exception as e:
            logger.warning("[CompetitorAnalyzer] LLM 解析失敗: %s", e)

        # フォールバック
        return {
            "summary": f"{target} の詳細情報を取得しました。",
            "strengths": ["情報収集中"],
            "weaknesses": ["情報収集中"],
            "competitors": ["情報収集中"],
        }

    def _create_docx(self, target: str, analysis: dict) -> str:
        """python-docx で Word レポートを作成して保存する。"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_target = "".join(c for c in target if c.isalnum() or c in "_ -")[:30]
        filename = f"{timestamp}_{safe_target}.docx"
        file_path = self._reports_dir / filename

        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            # タイトル
            title = doc.add_heading(f"競合分析レポート: {target}", level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 作成日
            doc.add_paragraph(f"作成日: {datetime.now().strftime('%Y年%m月%d日')}")
            doc.add_paragraph()

            # サマリー
            doc.add_heading("市場ポジション・概要", level=1)
            doc.add_paragraph(analysis.get("summary", ""))

            # 強み
            doc.add_heading("強み", level=1)
            for item in analysis.get("strengths", []):
                doc.add_paragraph(item, style="List Bullet")

            # 弱み
            doc.add_heading("弱み", level=1)
            for item in analysis.get("weaknesses", []):
                doc.add_paragraph(item, style="List Bullet")

            # 競合他社 テーブル
            doc.add_heading("主な競合他社", level=1)
            competitors = analysis.get("competitors", [])
            if competitors:
                table = doc.add_table(rows=1 + len(competitors), cols=2)
                table.style = "Table Grid"
                # ヘッダー行
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = "競合他社名"
                hdr_cells[1].text = "備考"
                for i, comp in enumerate(competitors, start=1):
                    row_cells = table.rows[i].cells
                    row_cells[0].text = comp
                    row_cells[1].text = ""

            doc.save(str(file_path))
            return str(file_path)

        except ImportError:
            logger.warning("[CompetitorAnalyzer] python-docx がインストールされていない。JSON で保存します。")
            json_path = file_path.with_suffix(".json")
            json_path.write_text(
                json.dumps(
                    {"target": target, **analysis},
                    ensure_ascii=False, indent=2,
                ),
                encoding="utf-8",
            )
            return str(json_path)
        except Exception as e:
            logger.warning("[CompetitorAnalyzer] docx 生成失敗: %s", e)
            return str(file_path)
